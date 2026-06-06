from __future__ import annotations

import io
import os
from datetime import datetime
from typing import Any, List

import pandas as pd
from fastapi import APIRouter, Depends, File, HTTPException, Query, UploadFile
from fastapi.responses import FileResponse, StreamingResponse
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill
from pydantic import BaseModel, Field
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
from sqlalchemy import func, or_, select
from sqlalchemy.orm import Session
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from app.db.database import get_db
from app.db.models import ResourceLibraryItemModel

# ── PDF 中文字体支持 ──────────────────────────────────────────
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
try:
    pdfmetrics.registerFont(TTFont("ChineseFont", "C:/Windows/Fonts/Deng.ttf"))
    CHINESE_FONT = "ChineseFont"
except Exception:
    CHINESE_FONT = "Helvetica"

router = APIRouter(prefix="/api/resource-library", tags=["资料库"])


def _parse_datetime(value: str | None) -> datetime | None:
    if not value:
        return None
    value = value.strip()
    for fmt in ("%Y-%m-%d %H:%M:%S", "%Y-%m-%dT%H:%M:%S", "%Y-%m-%d"):
        try:
            return datetime.strptime(value, fmt)
        except ValueError:
            continue
    raise ValueError(f"Invalid datetime: {value}")


def _format_datetime(value: datetime | None) -> str:
    if not value:
        return ""
    return value.strftime("%Y-%m-%d %H:%M:%S")


def _serialize(item: ResourceLibraryItemModel) -> dict[str, Any]:
    return {
        "id": item.id,
        "name": item.name,
        "netdiskType": item.netdisk_type,
        "url": item.url,
        "createdAt": _format_datetime(item.source_created_at),
        "updatedAt": _format_datetime(item.source_updated_at),
    }


class ResourceItemPayload(BaseModel):
    id: int | None = Field(default=None, ge=1)
    name: str = Field(..., min_length=1, max_length=512)
    netdiskType: str = Field("", max_length=64)
    url: str = Field(..., min_length=1)
    createdAt: str = ""
    updatedAt: str = ""


@router.get("")
def list_resources(
    keyword: str = Query("", max_length=200),
    netdiskType: str = Query("", max_length=64),
    page: int = Query(1, ge=1),
    pageSize: int = Query(20, ge=1, le=100),
    db: Session = Depends(get_db),
):
    stmt = select(ResourceLibraryItemModel)
    count_stmt = select(func.count()).select_from(ResourceLibraryItemModel)
    conditions = []

    if keyword.strip():
        pattern = f"%{keyword.strip()}%"
        conditions.append(
            or_(
                ResourceLibraryItemModel.name.like(pattern),
                ResourceLibraryItemModel.url.like(pattern),
            )
        )
    if netdiskType.strip():
        conditions.append(ResourceLibraryItemModel.netdisk_type == netdiskType.strip())

    for condition in conditions:
        stmt = stmt.where(condition)
        count_stmt = count_stmt.where(condition)

    total = db.scalar(count_stmt) or 0
    items = db.scalars(
        stmt.order_by(
            ResourceLibraryItemModel.source_updated_at.desc(),
            ResourceLibraryItemModel.id.desc(),
        )
        .offset((page - 1) * pageSize)
        .limit(pageSize)
    ).all()

    return {
        "success": True,
        "data": [_serialize(item) for item in items],
        "pagination": {
            "page": page,
            "pageSize": pageSize,
            "total": total,
            "totalPages": (total + pageSize - 1) // pageSize if total else 0,
        },
    }


@router.get("/netdisk-types")
def list_netdisk_types(db: Session = Depends(get_db)):
    rows = db.execute(
        select(ResourceLibraryItemModel.netdisk_type, func.count())
        .where(ResourceLibraryItemModel.netdisk_type != "")
        .group_by(ResourceLibraryItemModel.netdisk_type)
        .order_by(func.count().desc(), ResourceLibraryItemModel.netdisk_type.asc())
    ).all()
    return {
        "success": True,
        "data": [{"name": name, "count": count} for name, count in rows],
    }


@router.post("")
def create_resource(payload: ResourceItemPayload, db: Session = Depends(get_db)):
    # URL 去重检查
    existing_by_url = db.scalar(
        select(ResourceLibraryItemModel).where(
            ResourceLibraryItemModel.url == payload.url.strip()
        )
    )
    if existing_by_url:
        raise HTTPException(status_code=409, detail="网盘链接已存在")

    if payload.id is not None and db.get(ResourceLibraryItemModel, payload.id):
        raise HTTPException(status_code=409, detail="资源 ID 已存在")

    next_id = payload.id
    if next_id is None:
        next_id = (db.scalar(select(func.max(ResourceLibraryItemModel.id))) or 0) + 1

    item = ResourceLibraryItemModel(
        id=next_id,
        name=payload.name.strip(),
        netdisk_type=payload.netdiskType.strip(),
        url=payload.url.strip(),
        source_created_at=_parse_datetime(payload.createdAt) or datetime.now(),
        source_updated_at=_parse_datetime(payload.updatedAt) or datetime.now(),
    )
    db.add(item)
    db.flush()
    return {"success": True, "data": _serialize(item)}


@router.put("/{resource_id}")
def update_resource(resource_id: int, payload: ResourceItemPayload, db: Session = Depends(get_db)):
    item = db.get(ResourceLibraryItemModel, resource_id)
    if not item:
        raise HTTPException(status_code=404, detail="资源不存在")

    item.name = payload.name.strip()
    item.netdisk_type = payload.netdiskType.strip()
    item.url = payload.url.strip()
    item.source_created_at = _parse_datetime(payload.createdAt) or item.source_created_at
    item.source_updated_at = _parse_datetime(payload.updatedAt) or datetime.now()
    db.flush()
    return {"success": True, "data": _serialize(item)}


@router.delete("/{resource_id}")
def delete_resource(resource_id: int, db: Session = Depends(get_db)):
    item = db.get(ResourceLibraryItemModel, resource_id)
    if not item:
        raise HTTPException(status_code=404, detail="资源不存在")
    db.delete(item)
    return {"success": True, "message": "资源已删除"}


# ── 网盘类型自动识别映射（键名必须与数据库已有类型名称一致）──────
_NETDISK_KEYWORDS = {
    "夸克网盘": ["pan.quark.cn", "quark.cn"],
    "百度网盘": ["pan.baidu.com", "baidu.com"],
    "阿里网盘": ["alipan.com", "aliyundrive.com"],
    "迅雷云盘": ["pan.xunlei.com", "xunlei.com"],
    "蓝奏云": ["lanzou.com", "lanzous.com", "lanzoui.com"],
    "天翼云盘": ["cloud.189.cn", "189.cn"],
    "微云": ["weiyun.com"],
    "OneDrive": ["onedrive.live.com", "sharepoint.com"],
    "Google Drive": ["drive.google.com"],
    "Dropbox": ["dropbox.com"],
}

# 数据库中"其他"作为兜底类型
_NETDISK_FALLBACK = "其他"


def _detect_netdisk_type(url: str) -> str:
    url_lower = url.lower()
    for ntype, domains in _NETDISK_KEYWORDS.items():
        for d in domains:
            if d in url_lower:
                return ntype
    return _NETDISK_FALLBACK


# ── 导入 Excel ───────────────────────────────────────────────────────────
@router.post("/import")
async def import_resources(file: UploadFile = File(...), db: Session = Depends(get_db)):
    if not file.filename or not file.filename.endswith((".xlsx", ".xls")):
        raise HTTPException(status_code=400, detail="仅支持 .xlsx / .xls 文件")

    try:
        content = await file.read()
        df = pd.read_excel(io.BytesIO(content), header=None, usecols=[0, 1])
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Excel 解析失败：{e}")

    now_dt = datetime.now()
    success, dup_urls, errors = 0, [], []

    # 判断是否含表头：读取第一行，若第 0 列包含非资料名称文本则跳过
    if len(df) > 0:
        first_val = str(df.iloc[0, 0]).strip() if pd.notna(df.iloc[0, 0]) else ""
        if not first_val or not first_val.startswith("http"):
            # 看起来是表头行，跳过
            df = df.iloc[1:].reset_index(drop=True)

    # 计算导入起始 ID（正数累加，取当前最大 ID + 1）
    max_existing = db.scalar(select(func.max(ResourceLibraryItemModel.id))) or 0
    next_id = max_existing + 1

    for idx, row in df.iterrows():
        name = str(row[0]).strip() if pd.notna(row[0]) else ""
        url = str(row[1]).strip() if pd.notna(row[1]) else ""
        if not name or not url:
            continue
        if not url.startswith("http"):
            errors.append(f"第 {idx+1} 行：链接格式无效")
            continue

        existing = db.scalar(
            select(ResourceLibraryItemModel).where(ResourceLibraryItemModel.url == url)
        )
        if existing:
            dup_urls.append(url)
            continue

        item = ResourceLibraryItemModel(
            id=next_id,
            name=name,
            netdisk_type=_detect_netdisk_type(url),
            url=url,
            source_created_at=now_dt,
            source_updated_at=now_dt,
        )
        db.add(item)
        next_id += 1
        success += 1

    db.flush()
    return {
        "success": True,
        "data": {
            "imported": success,
            "duplicates": len(dup_urls),
            "errors": len(errors),
            "dup_urls": dup_urls[:10],
            "error_msgs": errors[:10],
        },
    }


# ── 导出 ────────────────────────────────────────────────────────────────
@router.get("/export")
def export_resources(
    keyword: str = Query("", max_length=200),
    netdiskType: str = Query("", max_length=64),
    limit: int = Query(0, ge=0, le=10000),
    export_format: str = Query("excel", max_length=10, alias="format"),
    ids_str: str = Query("", max_length=2000, alias="ids"),
    db: Session = Depends(get_db),
):
    stmt = select(ResourceLibraryItemModel)
    parsed_ids: list[int] = [int(x) for x in ids_str.split(",") if x.strip().isdigit()]
    if parsed_ids:
        stmt = stmt.where(ResourceLibraryItemModel.id.in_(parsed_ids))
    if keyword.strip():
        pat = f"%{keyword.strip()}%"
        stmt = stmt.where(
            or_(
                ResourceLibraryItemModel.name.like(pat),
                ResourceLibraryItemModel.url.like(pat),
            )
        )
    if netdiskType.strip():
        stmt = stmt.where(ResourceLibraryItemModel.netdisk_type == netdiskType.strip())

    items = db.scalars(
        stmt.order_by(
            ResourceLibraryItemModel.source_updated_at.desc(),
            ResourceLibraryItemModel.id.desc(),
        ).limit(limit if limit > 0 else None)
    ).all()

    if not items:
        raise HTTPException(status_code=404, detail="没有可导出的数据")

    export_dir = os.path.join(os.path.dirname(__file__), "..", "..", "exports")
    os.makedirs(export_dir, exist_ok=True)
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    fname = f"资料库导出_{ts}"

    if export_format == "pdf":
        path = _export_pdf(os.path.join(export_dir, f"{fname}.pdf"), items)
    elif export_format == "word":
        path = _export_word(os.path.join(export_dir, f"{fname}.docx"), items)
    else:
        path = _export_excel(os.path.join(export_dir, f"{fname}.xlsx"), items)

    return FileResponse(
        path=path,
        filename=os.path.basename(path),
        media_type="application/octet-stream",
    )


# now() removed - use datetime.now() directly


def _export_excel(path: str, items: list[ResourceLibraryItemModel]) -> str:
    wb = Workbook()
    ws = wb.active
    ws.title = "资料库"
    # 列顺序与导入兼容：前两列必须为 资料名称 / 网盘链接
    headers = ["资料名称", "网盘链接", "网盘类型", "创建时间", "更新时间"]
    ws.append(headers)
    for cell in ws[1]:
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill("solid", fgColor="D94E28")
        cell.alignment = Alignment(horizontal="center", vertical="center")
    for item in items:
        ws.append([
            item.name,
            item.url,
            item.netdisk_type or "未分类",
            _format_datetime(item.source_created_at),
            _format_datetime(item.source_updated_at),
        ])
    for col in ws.columns:
        max_len = max(len(str(c.value or "")) for c in col)
        ws.column_dimensions[col[0].column_letter].width = min(max_len + 4, 60)
    wb.save(path)
    return path


def _export_pdf(path: str, items: list[ResourceLibraryItemModel]) -> str:
    doc = SimpleDocTemplate(path, pagesize=A4,
                            rightMargin=36, leftMargin=36, topMargin=36, bottomMargin=36)
    styles = getSampleStyleSheet()
    # 所有文本使用中文字体
    styles["Normal"].fontName = CHINESE_FONT
    title_style = ParagraphStyle("Title", fontName=CHINESE_FONT, fontSize=16,
                                textColor=colors.HexColor("#D94E28"), spaceAfter=12, alignment=1)
    normal = styles["Normal"]
    table_data = [["ID", "资料名称", "网盘类型", "链接", "更新时间"]]
    for item in items:
        table_data.append([
            str(item.id),
            Paragraph(item.name, normal),
            Paragraph(item.netdisk_type or "未分类", normal),
            Paragraph(item.url, normal),
            _format_datetime(item.source_updated_at) or "",
        ])
    t = Table(table_data, colWidths=[30, 120, 60, 180, 70], repeatRows=1)
    style = [
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#D94E28")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#444444")),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("FONTNAME", (0, 0), (-1, 0), CHINESE_FONT),
    ]
    t.setStyle(TableStyle(style))
    elems = [Paragraph("资料库导出", title_style), Spacer(1, 12), t]
    doc.build(elems)
    return path


def _export_word(path: str, items: list[ResourceLibraryItemModel]) -> str:
    document = Document()
    style = document.styles["Normal"]
    style.font.name = "微软雅黑"
    style.font.size = Pt(10)
    heading = document.add_heading("资料库导出", 0)
    heading.alignment = WD_TABLE_ALIGNMENT.CENTER

    table = document.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    for i, h in enumerate(["ID", "资料名称", "网盘类型", "链接", "更新时间"]):
        cell = table.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True

    for item in items:
        row = table.add_row()
        row.cells[0].text = str(item.id)
        row.cells[1].text = item.name
        row.cells[2].text = item.netdisk_type or "未分类"
        row.cells[3].text = item.url
        row.cells[4].text = _format_datetime(item.source_updated_at) or ""

    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    document.save(path)
    return path
